import datetime
import os
import webbrowser
import docx
import easygui
import pandas as pd
import docx2txt


def createFolder(nameOfFile):
    folderName = "Result/" + nameOfFile + " (Created at " +  str(datetime.datetime.now().strftime("%d-%m-%y (Time %H-%M-%S)" + ")"))
    os.mkdir(folderName)
    return folderName

def openDonatePage():
    if("Donate Now.url" in os.listdir()):
        webbrowser.open("https://www.paypal.com/donate?hosted_button_id=3LYZG2R4962BN&source=url", new=0, autoraise=True)
        os.rename("Donate Now.url", "Support Us.url")

def getFileName(filePath):
    fileName = os.path.basename(filePath)
    return fileName

def generateWordDocument(document, tags, row, resultFolder, wordFile):
    numberOfTags = len(tags)
    doc = document
    for para in doc.paragraphs:
        for run in para.runs:
            for i in range(numberOfTags):
                if (type(row[i+1]) == pd._tslib.Timestamp):
                    run.text = run.text.replace(str(tags[i]).strip(), str(pd.to_datetime(row[i+1]).strftime("%d/%m/%Y")))
                elif (pd.isnull(row[i+1]) != True):
                    run.text = run.text.replace(str(tags[i]).strip(), str(row[i+1]))
                else:
                    run.text = run.text.replace(str(tags[i]).strip(), "")
    doc.save(resultFolder + "/" + row[1] + " - " + wordFile)
    return None


def getExcelFile():
    df = pd.read_excel("InformationForDocumentGenerator.xlsx")
    return df

def main():
    wordDocumentFilePath = easygui.fileopenbox("Choose your file", multiple=False)
    doc = docx.Document(wordDocumentFilePath)
    df = getExcelFile()
    tags = list(df.columns)
    folder = createFolder(getFileName(wordDocumentFilePath))
    for row in df.itertuples():
        generateWordDocument(doc, tags, row, folder, getFileName(wordDocumentFilePath))

    openDonatePage()






if __name__ == '__main__':
    print("Thank you for using Document Generator")
    main()